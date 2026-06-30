import os
import fitz
from docx import Document
from typing import List, Optional
from enum import Enum
from abc import ABC, abstractmethod
from dataclasses import dataclass
from ..models import Span, SpanStatus

import logging
import re

logger = logging.getLogger(__name__)

class RedactionStyle(Enum):
    BLACK_BOX = "black_box"
    SOLID_BLOCKS = "solid_blocks"
    REDACTED_TYPE = "redacted_type"

@dataclass
class ExportConfig:
    style: RedactionStyle = RedactionStyle.BLACK_BOX

class BaseExporter(ABC):
    @abstractmethod
    def export(self, input_path: str, output_path: str, spans: List[Span], config: ExportConfig) -> None:
        pass

class PdfExporter(BaseExporter):
    def export(self, input_path: str, output_path: str, spans: List[Span], config: ExportConfig) -> None:
        logger.info(f"Export started for {input_path}")
        try:
            doc = fitz.open(input_path)
        except Exception as e:
            raise ValueError(f"Failed to open PDF, it may be corrupted: {e}")
            
        if doc.is_encrypted:
            if not doc.authenticate(""):
                raise ValueError("Encrypted PDFs are not supported without a password.")
                
        redacted_spans = [s for s in spans if s.status == SpanStatus.REDACTED]
        
        for page in doc:
            words = page.get_text("words")
            if not words:
                continue
                
            full_text = ""
            char_to_word = []
            word_rects = []
            
            for w_idx, w in enumerate(words):
                text = w[4]
                rect = fitz.Rect(w[:4])
                word_str = text + " "
                for char_idx in range(len(word_str)):
                    char_to_word.append(w_idx)
                full_text += word_str
                word_rects.append(rect)

            for span in redacted_spans:
                if not span.text:
                    continue
                    
                # Normalize all whitespace (including newlines) to a single space
                # because `full_text` is built from space-separated words without newlines.
                normalized_text = re.sub(r'\s+', ' ', span.text.strip())
                escaped_span = re.escape(normalized_text)
                pattern = escaped_span.replace(r'\ ', r'\s+')
                
                try:
                    matches = list(re.finditer(pattern, full_text, re.IGNORECASE))
                except Exception as e:
                    logger.warning(f"Regex error for text {span.text}: {e}")
                    matches = []
                
                if not matches:
                    start_idx = full_text.find(normalized_text)
                    if start_idx != -1:
                        class SimpleMatch:
                            def __init__(self, start, end):
                                self._start = start
                                self._end = end
                            def start(self): return self._start
                            def end(self): return self._end
                        matches = [SimpleMatch(start_idx, start_idx + len(normalized_text))]
                
                if not matches:
                    logger.warning(f"Text not found: {span.text}")
                    continue
                    
                if len(matches) > 1:
                    logger.info(f"Multiple matches found for text: {span.text}")
                    
                for match in matches:
                    start_char = match.start()
                    end_char = match.end() - 1
                    
                    if start_char < 0 or end_char >= len(char_to_word):
                        logger.warning(f"Invalid coordinates for text: {span.text}")
                        continue
                        
                    start_word_idx = char_to_word[start_char]
                    end_word_idx = char_to_word[end_char]
                    
                    merged_rect = fitz.Rect(word_rects[start_word_idx])
                    for idx in range(start_word_idx + 1, end_word_idx + 1):
                        merged_rect.include_rect(word_rects[idx])
                        
                    if config.style == RedactionStyle.BLACK_BOX:
                        page.add_redact_annot(merged_rect, fill=(0, 0, 0))
                    elif config.style == RedactionStyle.SOLID_BLOCKS:
                        replacement = "█" * len(span.text)
                        page.add_redact_annot(merged_rect, text=replacement, fill=(0, 0, 0), text_color=(1, 1, 1))
                    elif config.style == RedactionStyle.REDACTED_TYPE:
                        replacement = f"[REDACTED:{span.type.value.upper()}]"
                        page.add_redact_annot(merged_rect, text=replacement, fill=(0, 0, 0), text_color=(1, 1, 1))
                        
                    logger.info(f"Coordinates located and redaction applied for entity: {span.type}")
                    
            page.apply_redactions()
            
        try:
            doc.save(output_path)
            logger.info("PDF saved successfully.")
        except Exception as e:
            raise ValueError(f"Failed to save redacted PDF: {e}")
        finally:
            doc.close()
            
        logger.info(f"Entities redacted: {len(redacted_spans)}")

class DocxExporter(BaseExporter):
    def export(self, input_path: str, output_path: str, spans: List[Span], config: ExportConfig) -> None:
        try:
            doc = Document(input_path)
        except Exception as e:
            raise ValueError(f"Failed to open DOCX: {e}")

        redacted_spans = [s for s in spans if s.status == SpanStatus.REDACTED]

        def get_replacement_text(span: Span) -> str:
            if config.style == RedactionStyle.BLACK_BOX:
                return "█" * len(span.text)  # fallback since docx doesn't draw true rectangles easily without shapes
            elif config.style == RedactionStyle.SOLID_BLOCKS:
                return "█" * len(span.text)
            else:
                return f"[REDACTED:{span.type.value.upper()}]"

        def redact_paragraph(p):
            for span in redacted_spans:
                old_text = span.text
                if not old_text or old_text not in p.text:
                    continue
                
                # Replace multiple occurrences within the same paragraph safely
                while old_text in p.text:
                    full_text = ""
                    char_to_run = []
                    
                    for run_idx, run in enumerate(p.runs):
                        for char_idx in range(len(run.text)):
                            char_to_run.append((run_idx, char_idx))
                        full_text += run.text
                    
                    start_idx = full_text.find(old_text)
                    if start_idx == -1:
                        break  # Safety break
                        
                    end_idx = start_idx + len(old_text) - 1
                    
                    start_run_idx, start_char_idx = char_to_run[start_idx]
                    end_run_idx, end_char_idx = char_to_run[end_idx]
                    
                    replacement = get_replacement_text(span)
                    
                    if start_run_idx == end_run_idx:
                        # Single run replacement (fully enclosed in one run)
                        run = p.runs[start_run_idx]
                        run.text = run.text[:start_char_idx] + replacement + run.text[end_char_idx + 1:]
                    else:
                        # Multi-run replacement (text spans across formatting boundaries)
                        start_run = p.runs[start_run_idx]
                        start_run.text = start_run.text[:start_char_idx] + replacement
                        
                        # Clear intermediate runs
                        for r_idx in range(start_run_idx + 1, end_run_idx):
                            p.runs[r_idx].text = ""
                            
                        # Trim the end run
                        end_run = p.runs[end_run_idx]
                        end_run.text = end_run.text[end_char_idx + 1:]

        # Redact main document paragraphs
        for p in doc.paragraphs:
            redact_paragraph(p)
            
        # Redact tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        redact_paragraph(p)
                        
        # Redact headers and footers
        for section in doc.sections:
            if section.header is not None:
                for p in section.header.paragraphs:
                    redact_paragraph(p)
                for table in section.header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                redact_paragraph(p)
            if section.footer is not None:
                for p in section.footer.paragraphs:
                    redact_paragraph(p)
                for table in section.footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                redact_paragraph(p)

        try:
            doc.save(output_path)
        except Exception as e:
            raise ValueError(f"Failed to save redacted DOCX: {e}")

class TxtExporter(BaseExporter):
    def export(self, input_path: str, output_path: str, spans: List[Span], config: ExportConfig) -> None:
        try:
            with open(input_path, "r", encoding="utf-8") as f:
                text = f.read()
        except Exception as e:
            raise ValueError(f"Failed to read TXT file: {e}")
            
        redacted_spans = [s for s in spans if s.status == SpanStatus.REDACTED]
        redacted_spans.sort(key=lambda s: s.start_index, reverse=True)
        
        for span in redacted_spans:
            if config.style == RedactionStyle.SOLID_BLOCKS or config.style == RedactionStyle.BLACK_BOX:
                placeholder = "█" * len(span.text)
            elif config.style == RedactionStyle.REDACTED_TYPE:
                placeholder = f"[REDACTED:{span.type.value.upper()}]"
            else:
                placeholder = "█" * len(span.text)
                
            text = text[:span.start_index] + placeholder + text[span.end_index:]
            
        try:
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(text)
        except Exception as e:
            raise ValueError(f"Failed to save TXT file: {e}")

class ExportService:
    @classmethod
    def generate_export(cls, input_path: str, mime_type: str, spans: List[Span], style: str = "black_box") -> str:
        """
        Detect file type and apply the appropriate redaction exporter.
        Returns the output file path.
        """
        output_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "exports")
        os.makedirs(output_dir, exist_ok=True)
        
        base_name = os.path.basename(input_path)
        name, _ = os.path.splitext(base_name)
        
        try:
            config_style = RedactionStyle(style)
        except ValueError:
            config_style = RedactionStyle.BLACK_BOX
            
        config = ExportConfig(style=config_style)
        
        if "pdf" in mime_type or input_path.lower().endswith(".pdf"):
            out_name = f"{name}_redacted.pdf"
            out_path = os.path.join(output_dir, out_name)
            exporter = PdfExporter()
            exporter.export(input_path, out_path, spans, config)
            return out_path
            
        elif "wordprocessingml" in mime_type or input_path.lower().endswith(".docx"):
            out_name = f"{name}_redacted.docx"
            out_path = os.path.join(output_dir, out_name)
            exporter = DocxExporter()
            exporter.export(input_path, out_path, spans, config)
            return out_path
            
        else:
            out_name = f"{name}_redacted.txt"
            out_path = os.path.join(output_dir, out_name)
            exporter = TxtExporter()
            exporter.export(input_path, out_path, spans, config)
            return out_path
